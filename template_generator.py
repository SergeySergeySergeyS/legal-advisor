"""
Модуль генерации юридических документов (претензий, жалоб, исков, ходатайств)
ФИНАЛЬНАЯ версия со всеми улучшениями:
- Универсальные правила подсудности (ст. 23, 24, 30 ГПК РФ)
- Лимит 100 000 руб. для мировых судей
- Исключительная подсудность для недвижимости
- Правильная подсудность для деликтов (по месту жительства ответчика)
- 6 категорий с валидацией законов
- Два сценария наследственных споров (ст. 1153, 1155 ГК РФ)
- ТСТ (Постановление № 924)
- Условная госпошлина: освобождение только для ЗоЗПП, трудовых, ГИБДД, алиментов
- Условный расчёт цены иска: только для имущественных исков
- Правильная терминология: "моральный вред" (не "моральный ущерб")
- ЗАПРЕТ на моральный вред в: наследстве, выселении, деликтах (затопление)
- Расширенный фильтр мусора в приложениях
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import re
from gigachat import GigaChat
from gigachat.models import Chat, Messages, MessagesRole


def safe_decode(content):
    """Безопасно декодирует байты в строку"""
    if isinstance(content, bytes):
        return content.decode('utf-8', errors='ignore')
    return content


# === ДОПУСТИМЫЕ ЗАКОНЫ ПО КАТЕГОРИЯМ ===
VALID_LAWS_BY_CATEGORY = {
    "🛒 Защита прав потребителей": [
        "о защите прав потребителей", "зозпп", "2300-1",
        "гражданский кодекс", "гк рф", "потребител",
        "налоговый кодекс", "нк рф", "333.36",
        "постановление правительства", "№ 924", "№924",
        "технически сложн", "гпк рф", "подсудн"
    ],
    "💼 Трудовые споры": [
        "трудовой кодекс", "тк рф", "коллективн", "заработн",
        "увольнен", "дисциплинарн", "гпк рф", "нк рф", "333.36"
    ],
    "🚗 Споры с ГИБДД": [
        "коап", "коап рф", "правила дорожного", "пдд",
        "безопасност.*дорож", "осаго", "каско",
        "транспортн", "административн", "гпк рф"
    ],
    "👨‍👩‍👧 Семейное право": [
        "семейный кодекс", "ск рф", "алимент", "развод",
        "брачн", "раздел имуществ", "усыновлен", "гпк рф", "подсудн",
        "налоговый кодекс", "нк рф", "333.19", "госпошлин"
    ],
    "📜 Наследство": [
        "гражданский кодекс", "гк рф", "наследств", "завещан",
        "наследодател", "обязательн.*доля", "гпк рф",
        "налоговый кодекс", "нк рф", "333.19", "госпошлин"
    ],
    "🏠 Жилищные вопросы (вкл. ЖКХ)": [
        "жилищный кодекс", "жк рф",
        "постановление правительства", "№ 354", "№ 491", "№354", "№491",
        "гражданский кодекс", "гк рф",
        "управляющ", "жиль", "приватизац", "гпк рф",
        "налоговый кодекс", "нк рф", "333.19", "госпошлин"
    ]
}

# === ЗАПРЕЩЁННЫЕ ЗАКОНЫ ПО КАТЕГОРИЯМ ===
INVALID_LAWS_BY_CATEGORY = {
    "🛒 Защита прав потребителей": ["тк рф", "коап", "жилищн", "ск рф"],
    "💼 Трудовые споры": ["о защите прав потребителей", "коап", "жилищн", "ск рф"],
    "🚗 Споры с ГИБДД": ["о защите прав потребителей", "тк рф", "жилищн", "ск рф", "гк рф"],
    "👨‍👩‍👧 Семейное право": ["тк рф", "коап", "жилищн", "зозпп"],
    "📜 Наследство": ["тк рф", "коап", "жилищн", "зозпп", "ск рф"],
    "🏠 Жилищные вопросы (вкл. ЖКХ)": ["о защите прав потребителей", "тк рф", "коап", "ск рф"]
}

# === ТЕХНИЧЕСКИ СЛОЖНЫЕ ТОВАРЫ ===
TST_EXAMPLES = """
К технически сложным товарам относятся (Перечень, утв. Постановлением Правительства РФ от 10.11.2011 № 924):
- смартфоны, мобильные телефоны
- компьютеры (стационарные, портативные, планшеты)
- телевизоры, мониторы
- автомобили, мотоциклы
- бытовая техника (стиральные машины, холодильники, посудомойки)
- фото- и видеокамеры
- часы (наручные, карманные — электронные/механические с 2+ функциями)
- садовая техника (газонокосилки, бензопилы)
- навигационное оборудование
"""


# === ПРОМТЫ ДЛЯ РАЗНЫХ ТИПОВ ДОКУМЕНТОВ ===
DOCUMENT_PROMPTS = {
    "претензия": """Ты — юридический помощник. Создай ПРЕТЕНЗИЮ к контрагенту.

КРИТИЧЕСКИ ВАЖНО:
- Используй ТОЛЬКО информацию из ПОСЛЕДНЕГО вопроса клиента
- НЕ выдумывай факты
- Если данные не указаны — оставь поле ПУСТЫМ
- НЕ используй никакие скобки: ни круглые (), ни квадратные [], ни угловые <>
- Используй термин "моральный вред" (НЕ "моральный ущерб")

ВЕРНИ РЕЗУЛЬТАТ СТРОГО В ТАКОМ ФОРМАТЕ:

ТИП: претензия
КОМУ: [должность и наименование организации]
ОТ: [ФИО, адрес, телефон — если нет, оставь пустым]
СИТУАЦИЯ: [4-6 предложений от ПЕРВОГО ЛИЦА]
ЗАКОНЫ: [статьи законов через точку с запятой]
ТРЕБОВАНИЯ: [нумерованный список]
ПРИЛОЖЕНИЯ: [нумерованный список]

ВАЖНО:
- НЕ используй markdown
- НЕ используй скобки любого типа""",

    "жалоба": """Ты — юридический помощник. Создай ЖАЛОБУ в государственный орган.

КРИТИЧЕСКИ ВАЖНО:
- Используй ТОЛЬКО информацию из ПОСЛЕДНЕГО вопроса клиента
- НЕ выдумывай факты
- Если данные не указаны — оставь поле ПУСТЫМ
- НЕ используй никакие скобки: ни круглые (), ни квадратные [], ни угловые <>
- Используй термин "моральный вред" (НЕ "моральный ущерб")

ВЕРНИ РЕЗУЛЬТАТ СТРОГО В ТАКОМ ФОРМАТЕ:

ТИП: жалоба
КОМУ: [наименование госоргана]
ОТ: [ФИО, адрес, телефон — если нет, оставь пустым]
СИТУАЦИЯ: [4-6 предложений от ПЕРВОГО ЛИЦА]
ЗАКОНЫ: [статьи законов через точку с запятой]
ТРЕБОВАНИЯ: [нумерованный список — ПРОВЕСТИ ПРОВЕРКУ, ПРИНЯТЬ МЕРЫ]
ПРИЛОЖЕНИЯ: [нумерованный список]

ВАЖНО:
- НЕ используй markdown
- НЕ используй скобки любого типа""",

    "иск": """Ты — юридический помощник. Создай ИСКОВОЕ ЗАЯВЛЕНИЕ в суд.

КРИТИЧЕСКИ ВАЖНО:
- Используй ТОЛЬКО информацию из ПОСЛЕДНЕГО вопроса клиента
- НЕ выдумывай факты
- Если конкретные данные (ФИО, даты, суммы) не указаны, используй ТОЛЬКО прочерки: ___________.
- Пиши текст от ПЕРВОГО ЛИЦА ("Я, ___________, прошу...").
- Используй термин "моральный вред" (НЕ "моральный ущерб")

ВЕРНИ РЕЗУЛЬТАТ СТРОГО В ТАКОМ ФОРМАТЕ:

ТИП: исковое заявление
КОМУ: [Определи подсудность САМ по правилам ниже]
ИСТЕЦ: ___________ (ФИО, адрес, телефон)
ОТВЕТЧИК: ___________ (ФИО или наименование, адрес)
СИТУАЦИЯ: [4-6 предложений от ПЕРВОГО ЛИЦА]
ЗАКОНЫ: [статьи законов через точку с запятой]
ЦЕНА_ИСКА: [Расчёт цены иска ИЛИ фраза "Иск неимущественного характера"]
ТРЕБОВАНИЯ: [нумерованный список]
ПРИЛОЖЕНИЯ: [нумерованный список]

=== УНИВЕРСАЛЬНЫЕ ПРАВИЛА ПОДСУДНОСТИ (ст. 23, 24, 30 ГПК РФ) ===

Определяй подсудность по следующей логике:

1. ТРУДОВЫЕ СПОРЫ → ВСЕГДА районный суд (ст. 24, 391 ГПК РФ).

2. СПОРЫ С ГИБДД → ВСЕГДА районный суд (ст. 30.1 КоАП РФ).

3. СЕМЕЙНЫЕ СПОРЫ:
   - Если ЕСТЬ спор о детях → ВСЕГДА районный суд (ст. 24 ГПК РФ).
   - Если НЕТ спора о детях и цена иска ДО 100 000 руб. → Мировой судья.
   - Если НЕТ спора о детях и цена иска СВЫШЕ 100 000 руб. → Районный суд.

4. ЗАЩИТА ПРАВ ПОТРЕБИТЕЛЕЙ:
   - Цена иска ДО 100 000 руб. → Мировой судья.
   - Цена иска СВЫШЕ 100 000 руб. → Районный суд.
   - По ст. 17 ЗоЗПП потребитель вправе подать иск ПО МЕСТУ СВОЕГО ЖИТЕЛЬСТВА (альтернативная подсудность).

5. НАСЛЕДСТВЕННЫЕ СПОРЫ И СПОРЫ О НЕДВИЖИМОСТИ (КРИТИЧЕСКИ ВАЖНО!):
   - ЕСЛИ спор касается прав на недвижимость (квартира, дом, земля) → ВСЕГДА Районный суд ПО МЕСТУ НАХОЖДЕНИЯ этого недвижимого имущества (Исключительная подсудность, ст. 30 ГПК РФ).
   - НИКОГДА не направляй иски о недвижимости к мировому судье.
   - НИКОГДА не используй "по месту жительства истца" для споров о недвижимости.
   - Если спор только о движимом имуществе до 100 000 руб. → Мировой судья.
   - Если спор о движимом имуществе свыше 100 000 руб. → Районный суд.

6. ЖИЛИЩНЫЕ СПОРЫ:
   - Споры о выселении, признании утратившим право пользования → ВСЕГДА районный суд ПО МЕСТУ НАХОЖДЕНИЯ недвижимости (ст. 30 ГПК РФ).
   - ДЕЛИКТНЫЕ ИСКИ (затопление соседями, причинение вреда имуществу) → ПО МЕСТУ ЖИТЕЛЬСТВА ОТВЕТЧИКА (ст. 28 ГПК РФ), НЕ по месту жительства истца!
   - Имущественные споры (перерасчёт, возмещение ущерба от УК) до 100 000 руб. → Мировой судья.
   - Свыше 100 000 руб. → Районный суд.

ФОРМАТ ЗАПИСИ В ПОЛЕ "КОМУ":
- Для мирового судьи: "Мировому судье судебного участка № ___ [район] г. [город]"
- Для районного суда: "В [название] районный суд г. [город]"
- Добавь указание на подсудность в зависимости от ситуации.

=== ПРАВИЛА РАСЧЁТА ЦЕНЫ ИСКА ===

Для ИМУЩЕСТВЕННЫХ исков:
- Защита прав потребителей: стоимость товара + неустойка (1% за день) + компенсация морального вреда
- Семейные споры (алименты): сумма в месяц × 12 месяцев
- Семейные споры (раздел имущества): стоимость доли истца
- Наследство (Сценарий Б — фактическое принятие): стоимость наследуемого имущества
- Жилищные споры (затопление, перерасчёт): сумма ущерба

Для НЕИМУЩЕСТВЕННЫХ исков пиши просто: "Иск неимущественного характера".
К неимущественным относятся:
- Иски о выселении / признании утратившим право пользования
- Иски о восстановлении срока принятия наследства (Сценарий А)
- Иски об установлении факта принятия наследства
- Иски об определении места жительства ребёнка
- Иски о расторжении брака (без раздела имущества)

=== ЗАПРЕТ НА МОРАЛЬНЫЙ ВРЕД В ОПРЕДЕЛЁННЫХ КАТЕГОРИЯХ ИСКОВ ===

КРИТИЧЕСКИ ВАЖНО: НЕ включай требование о компенсации морального вреда в следующие категории исков:
- Иски о выселении / признании утратившим право пользования жилым помещением (ст. 31, 35 ЖК РФ)
- Иски о восстановлении срока принятия наследства (ст. 1155 ГК РФ)
- Иски об установлении факта принятия наследства (ст. 1153 ГК РФ)
- Иски об определении места жительства ребёнка
- Иски о расторжении брака (без раздела имущества)
- Иски о возмещении ущерба от затопления квартиры (ст. 1064 ГК РФ) — деликтные иски
- Любые другие деликтные иски о возмещении имущественного ущерба

Причина: ст. 151 ГК РФ предусматривает компенсацию морального вреда только при нарушении личных неимущественных прав или когда это прямо предусмотрено законом.

МОРАЛЬНЫЙ ВРЕД можно взыскивать ТОЛЬКО в:
- Исках по ЗоЗПП (ст. 15 ЗоЗПП — прямо предусмотрено законом)
- Исках о возмещении вреда здоровью (ст. 1099 ГК РФ)
- Исках о защите чести, достоинства и деловой репутации (ст. 152 ГК РФ)
- Исках о незаконном увольнении (ст. 237 ТК РФ)
- Исках о разделе имущества (вместе с имущественными требованиями)

=== ПРАВИЛА ГОСПОШЛИНЫ ===

Категории, ОСВОБОЖДЁННЫЕ от госпошлины (НЕ включай квитанцию в приложения!):
- Защита прав потребителей (ст. 333.36 НК РФ, цена иска до 1 млн руб.)
- Трудовые споры (ст. 333.36 п.1 пп.1 НК РФ)
- Обжалование постановлений ГИБДД
- Взыскание алиментов (ст. 333.36 п.15 НК РФ)

Категории, где госпошлина ПЛАТИТСЯ (ОБЯЗАТЕЛЬНО включи квитанцию в приложения!):
- Семейные споры: развод (650 руб.), раздел имущества (% от цены иска)
- Наследственные споры
- Жилищные споры (включая неимущественные — 300 руб. для физлиц)
- Деликтные иски (затопление и т.д.)

=== ТЕХНИЧЕСКИ СЛОЖНЫЕ ТОВАРЫ (ТСТ) ===

""" + TST_EXAMPLES + """

ПРАВИЛА ДЛЯ ТСТ (ст. 18 п. 1 ЗоЗПП):
- В течение 15 дней — потребитель вправе отказаться от товара
- ПОСЛЕ 15 дней — только при существенном недостатке
- В правовое обоснование добавь: "Постановление Правительства РФ от 10.11.2011 № 924"

=== ПРАВИЛА ДЛЯ НАСЛЕДСТВЕННЫХ СПОРОВ ===

- СЦЕНАРИЙ А (Срок пропущен + НЕ было фактического принятия): 
  Требования: "1. Восстановить срок для принятия наследства. 2. Признать истца принявшим наследство" (ст. 1155 ГК РФ).
  Цена иска: "Иск неимущественного характера".
  БЕЗ морального вреда!
- СЦЕНАРИЙ Б (Было ФАКТИЧЕСКОЕ принятие): 
  Требования: "1. Установить факт принятия наследства. 2. Признать право собственности" (п. 2 ст. 1153 ГК РФ).
  Цена иска: стоимость наследуемого имущества.
  БЕЗ морального вреда!
- КВИТАНЦИЯ О ГОСПОШЛИНЕ ОБЯЗАТЕЛЬНА для наследственных дел!

=== ТРЕБОВАНИЯ ===

Для разных категорий:
- Защита прав потребителей: стоимость товара, неустойка, компенсация морального вреда, штраф 50%
- Семейные споры: алименты, раздел имущества (моральный вред — только при сопутствующих обстоятельствах)
- Наследство: зависит от сценария, БЕЗ морального вреда
- Жилищные споры: возмещение ущерба, перерасчёт, признание утратившим право (БЕЗ морального вреда для выселения и затопления!)
- Трудовые: восстановление на работе, средний заработок, моральный вред

НЕ включай пояснения типа "не заявлено", "отсутствует".

=== ПРИЛОЖЕНИЯ ===

Всегда включай:
1. Копия искового заявления для ответчика
2. Документы, подтверждающие обстоятельства дела
3. Расчёт цены иска (ТОЛЬКО для имущественных исков!)
4. Квитанция об уплате госпошлины (КРОМЕ исков по ЗоЗПП, трудовых, обжалования ГИБДД, алиментов)

НЕ включай:
- Копии статей законодательства
- Любые куски статей законов
- Устав ООО, свидетельство о регистрации

ВАЖНО:
- НЕ используй скобки любого типа
- НЕ используй placeholder'ы, только прочерки ___________
- НЕ используй markdown""",

    "ходатайство": """Ты — юридический помощник. Создай ХОДАТАЙСТВО.

КРИТИЧЕСКИ ВАЖНО:
- Используй ТОЛЬКО информацию из ПОСЛЕДНЕГО вопроса клиента
- НЕ выдумывай факты
- Если данные не указаны — оставь поле ПУСТЫМ
- НЕ используй никакие скобки
- Используй термин "моральный вред" (НЕ "моральный ущерб")

ВЕРНИ РЕЗУЛЬТАТ СТРОГО В ТАКОМ ФОРМАТЕ:

ТИП: ходатайство
КОМУ: [наименование органа]
ОТ: [ФИО, адрес, телефон — если нет, оставь пустым]
СИТУАЦИЯ: [3-5 предложений от ПЕРВОГО ЛИЦА]
ЗАКОНЫ: [статьи законов через точку с запятой]
ТРЕБОВАНИЯ: [нумерованный список — ПРОШУ: 1. Предоставить...]
ПРИЛОЖЕНИЯ: [нумерованный список]

ВАЖНО:
- НЕ используй markdown
- НЕ используй скобки любого типа"""
}


def analyze_chat_for_template(chat_history: list, category: str, auth_key: str, doc_type: str = "претензия") -> dict:
    """Анализирует историю чата и извлекает данные для документа"""
    try:
        prompt_template = DOCUMENT_PROMPTS.get(doc_type, DOCUMENT_PROMPTS["претензия"])
        
        last_user_question = ""
        last_assistant_answer = ""
        
        for msg in reversed(chat_history):
            if msg["role"] == "user" and not last_user_question:
                last_user_question = msg['content']
            elif msg["role"] == "assistant" and not last_assistant_answer:
                last_assistant_answer = msg['content']
            if last_user_question and last_assistant_answer:
                break
        
        prompt = f"""КАТЕГОРИЯ: {category}

ПОСЛЕДНИЙ ВОПРОС КЛИЕНТА:
{last_user_question}

ОТВЕТ ЮРИСТА:
{last_assistant_answer}

{prompt_template}

Создавай документ ТОЛЬКО на основе ПОСЛЕДНЕГО вопроса."""
        
        with GigaChat(
            credentials=auth_key,
            scope="GIGACHAT_API_PERS",
            verify_ssl_certs=False
        ) as giga:
            messages = [Messages(role=MessagesRole.USER, content=prompt)]
            response = giga.chat(Chat(messages=messages))
            result_text = safe_decode(response.choices[0].message.content)
        
        print(f"🔍 Сырой ответ ИИ для типа '{doc_type}':\n{result_text}\n")
        
        result = parse_template_text(result_text, doc_type)
        result['legal_basis'] = validate_legal_basis(result.get('legal_basis', ''), category)
        
        return result
        
    except Exception as e:
        print(f"❌ Ошибка анализа: {e}")
        return get_default_template(doc_type)


def validate_legal_basis(legal_basis: str, category: str) -> str:
    """Проверяет правовое обоснование на соответствие категории"""
    if not legal_basis or legal_basis == '[Статьи законов]':
        return legal_basis
    
    valid_keywords = VALID_LAWS_BY_CATEGORY.get(category, [])
    invalid_keywords = INVALID_LAWS_BY_CATEGORY.get(category, [])
    
    articles = [a.strip() for a in legal_basis.split(';') if a.strip()]
    
    filtered_articles = []
    for article in articles:
        article_lower = article.lower()
        is_invalid = any(kw in article_lower for kw in invalid_keywords)
        is_valid = any(re.search(kw, article_lower) for kw in valid_keywords) if valid_keywords else True
        
        if not is_invalid:
            filtered_articles.append(article)
        else:
            print(f"🚫 Удалена нерелевантная статья: {article}")
    
    if not filtered_articles:
        return legal_basis
    
    return '; '.join(filtered_articles)


def parse_template_text(text: str, doc_type: str = "претензия") -> dict:
    """Парсит текстовый ответ в словарь"""
    result = get_default_template(doc_type)
    
    try:
        lines = text.split('\n')
        current_section = None
        current_content = []
        
        def save_section():
            nonlocal current_section, current_content
            if current_section and current_content:
                content = '\n'.join(current_content).strip()
                if content:
                    result[current_section] = content
            current_content = []
        
        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                continue
            
            line_upper = line_stripped.upper()
            
            if any(x in line_upper for x in ['ТИП:', 'ТИП ДОКУМЕНТА:']):
                save_section()
                current_section = 'document_type'
                content = re.sub(r'^ТИП( ДОКУМЕНТА)?:\s*', '', line_stripped, flags=re.IGNORECASE)
                current_content = [content]
            
            elif any(x in line_upper for x in ['КОМУ:', 'АДРЕСАТ:', 'АДРЕСАТУ:', 'В СУД:', 'В:']):
                save_section()
                current_section = 'recipient'
                content = re.sub(r'^(КОМУ|АДРЕСАТ|АДРЕСАТУ|В СУД|В):\s*', '', line_stripped, flags=re.IGNORECASE)
                current_content = [content]
            
            elif any(x in line_upper for x in ['ИСТЕЦ:', 'ЗАЯВИТЕЛЬ:']):
                save_section()
                current_section = 'sender'
                content = re.sub(r'^(ИСТЕЦ|ЗАЯВИТЕЛЬ):\s*', '', line_stripped, flags=re.IGNORECASE)
                current_content = [content]
            
            elif any(x in line_upper for x in ['ОТ:', 'ОТ КОГО:']):
                save_section()
                current_section = 'sender'
                content = re.sub(r'^ОТ( КОГО)?:\s*', '', line_stripped, flags=re.IGNORECASE)
                current_content = [content]
            
            elif any(x in line_upper for x in ['ОТВЕТЧИК:']):
                save_section()
                current_section = 'defendant'
                content = re.sub(r'^ОТВЕТЧИК:\s*', '', line_stripped, flags=re.IGNORECASE)
                current_content = [content]
            
            elif any(x in line_upper for x in ['СИТУАЦИЯ:', 'ОПИСАНИЕ:', 'ФАКТЫ:', 'ОБСТОЯТЕЛЬСТВА:']):
                save_section()
                current_section = 'situation'
                content = re.sub(r'^(СИТУАЦИЯ|ОПИСАНИЕ|ФАКТЫ|ОБСТОЯТЕЛЬСТВА):\s*', '', line_stripped, flags=re.IGNORECASE)
                current_content = [content]
            
            elif any(x in line_upper for x in ['ЗАКОНЫ:', 'ПРАВОВОЕ ОБОСНОВАНИЕ', 'ПРАВОВАЯ БАЗА', 
                                                'СТАТЬИ:', 'НОРМАТИВНАЯ БАЗА', 'ОБОСНОВАНИЕ:']):
                save_section()
                current_section = 'legal_basis'
                content = re.sub(r'^(ЗАКОНЫ|ПРАВОВОЕ ОБОСНОВАНИЕ|ПРАВОВАЯ БАЗА|СТАТЬИ|НОРМАТИВНАЯ БАЗА|ОБОСНОВАНИЕ):\s*', 
                                '', line_stripped, flags=re.IGNORECASE)
                current_content = [content]
            
            elif any(x in line_upper for x in ['ЦЕНА ИСКА:', 'ЦЕНА_ИСКА:', 'СУММА ИСКА:']):
                save_section()
                current_section = 'claim_amount'
                content = re.sub(r'^(ЦЕНА ИСКА|ЦЕНА_ИСКА|СУММА ИСКА):\s*', '', line_stripped, flags=re.IGNORECASE)
                current_content = [content]
            
            elif any(x in line_upper for x in ['ТРЕБОВАНИЯ:', 'ТРЕБУЮ:', 'ПРОШУ:', 'ПРОШУ СУД:', 'ХОДАТАЙСТВУЮ:']):
                save_section()
                current_section = 'requirements'
                content = re.sub(r'^(ТРЕБОВАНИЯ|ТРЕБУЮ|ПРОШУ СУД|ПРОШУ|ХОДАТАЙСТВУЮ):\s*', '', line_stripped, flags=re.IGNORECASE)
                current_content = [content]
            
            elif any(x in line_upper for x in ['ПРИЛОЖЕНИЯ:', 'ПРИЛАГАЮ:', 'ПРИЛОЖЕННЫЕ ДОКУМЕНТЫ']):
                save_section()
                current_section = 'attachments'
                content = re.sub(r'^(ПРИЛОЖЕНИЯ|ПРИЛАГАЮ|ПРИЛОЖЕННЫЕ ДОКУМЕНТЫ):\s*', '', line_stripped, flags=re.IGNORECASE)
                current_content = [content]
            
            elif current_section:
                current_content.append(line_stripped)
        
        save_section()
        
        for key in result:
            if result[key]:
                result[key] = result[key].replace('**', '').replace('*', '')
                result[key] = result[key].replace('__', '').replace('_', '')
                result[key] = re.sub(r'\s+', ' ', result[key]).strip()
        
        print(f"✅ Распознанные поля для '{doc_type}': {list(result.keys())}")
        
    except Exception as e:
        print(f"❌ Ошибка парсинга: {e}")
    
    return result


def get_default_template(doc_type: str = "претензия") -> dict:
    """Возвращает шаблон по умолчанию"""
    return {
        "document_type": doc_type.upper(),
        "recipient": "",
        "sender": "",
        "situation": "",
        "legal_basis": "",
        "requirements": "",
        "attachments": "",
        "defendant": "",
        "claim_amount": ""
    }


def clean_template_hints(text: str) -> str:
    """Удаляет шаблонные подсказки, скобки и лишние символы"""
    if not text:
        return text
    
    patterns = [
        r'\(укажите[^)]*\)',
        r'\(если[^)]*\)',
        r'\(оставь[^)]*\)',
        r'\(если известен[^)]*\)',
        r'\(если имеются[^)]*\)',
        r'\(если понесены[^)]*\)',
        r'\(если такие[^)]*\)',
        r'\(при наличии[^)]*\)',
        r'\(при их наличии[^)]*\)',
        r'\(если таковые[^)]*\)',
        r'\(укажи[^)]*\)',
        r'\(если известно[^)]*\)',
        r'\(если они[^)]*\)',
        r'\(если есть[^)]*\)',
    ]
    
    for pattern in patterns:
        text = re.sub(pattern, '', text, flags=re.IGNORECASE)
    
    text = re.sub(r'\[([^\]]*)\]', r'\1', text)
    text = re.sub(r'<([^>]*)>', r'\1', text)
    text = re.sub(r'\([А-ЯA-Z][А-Яа-яA-Za-z\s,]+\)', '', text)
    
    text = text.strip()
    text = text.strip('"').strip("'").strip('«').strip('»')
    text = text.strip()
    
    text = re.sub(r'\s+', ' ', text).strip()
    text = text.rstrip(',;')
    
    return text


def format_claim_amount(text: str) -> str:
    """Форматирует цену иска — разбивает на строки"""
    if not text:
        return text
    
    if ' - ' in text and '\n' not in text:
        parts = [p.strip() for p in text.split(' - ') if p.strip()]
        if len(parts) > 1:
            return '\n'.join(parts)
    
    return text


def remove_duplicate_headers(text: str) -> str:
    """Удаляет дублирующиеся заголовки"""
    if not text:
        return text
    
    text = re.sub(r'Цена иска:\s*Цена иска:', 'Цена иска:', text, flags=re.IGNORECASE)
    text = re.sub(r'^\s*ПРОШУ СУД:\s*', '', text, flags=re.IGNORECASE)
    text = re.sub(r'^\s*1\.\s*ПРОШУ СУД:\s*', '', text, flags=re.IGNORECASE)
    
    return text


def is_garbage_attachment(line: str) -> bool:
    """Проверяет, является ли приложение мусором"""
    line_lower = line.lower().strip()
    line_cleaned = re.sub(r'^\d+[\.\)]\s*', '', line_lower).strip()
    
    garbage_patterns = [
        r'^\d+\s*(налогов|нк\s*рф)',
        r'^\d+\s+(гк|тк|коап|жк|закон|кодекс)',
        r'^(налогов|нк\s*рф)',
        r'^(гк|тк|коап|жк|закон|кодекс)\s+рф',
        r'налогов\s+кодекс',
        r'стать(и|я|ей)\s+законодательств',
        r'копи(я|и|ю)\s+статей',
        r'копи(я|и|ю)\s+законодательств',
        r'устав\s+(ооо|зао|ао|пао)',
        r'свидетельств[ао]\s+о\s+регистрац',
        r'егрюл|егрип',
        r'^\d{1,4}\s+(гк|тк|коап|жк|закон|кодекс|ск)',
        r'ст\.?\s*\d+',
        r'пункт\s+\d+',
        r'333\.36',
        r'1153\s*гк',
        r'1154\s*гк',
        r'1155\s*гк',
    ]
    
    for pattern in garbage_patterns:
        if re.search(pattern, line_lower) or re.search(pattern, line_cleaned):
            return True
    
    return False


def is_non_property_claim(template_data: dict) -> bool:
    """Определяет, является ли иск неимущественным"""
    claim_amount = (template_data.get('claim_amount', '') or '').lower()
    if 'неимущественн' in claim_amount:
        return True
    
    requirements = (template_data.get('requirements', '') or '').lower()
    
    non_property_markers = [
        'выселить',
        'утративш',
        'восстановить срок',
        'установить факт',
        'место жительства ребенка',
        'место жительства ребёнка',
        'расторгнуть брак',
        'порядок общения',
        'лишить родительских',
        'признать недостойн',
        'признать безвестно отсутствующ',
        'признать умерш',
    ]
    
    property_markers = [
        'взыскать',
        'стоимость',
        'раздел имуществ',
        'признать право собственности',
        'возместить ущерб',
        'неустойк',
    ]
    
    has_property = any(marker in requirements for marker in property_markers)
    if has_property:
        return False
    
    has_non_property = any(marker in requirements for marker in non_property_markers)
    if has_non_property:
        return True
    
    return False


def is_fee_exempt_category(template_data: dict) -> bool:
    """Определяет, освобождён ли иск от госпошлины"""
    legal_basis = (template_data.get('legal_basis', '') or '').lower()
    requirements = (template_data.get('requirements', '') or '').lower()
    situation = (template_data.get('situation', '') or '').lower()
    
    combined = f"{legal_basis} {requirements} {situation}"
    
    if any(x in combined for x in ['зозпп', 'о защите прав потребителей', '2300-1']):
        return True
    
    if any(x in combined for x in ['тк рф', 'трудов', 'восстановить на работе', 'вынужденн.*прогул']):
        return True
    
    if 'алимент' in combined and 'раздел имуществ' not in combined:
        return True
    
    if any(x in combined for x in ['коап', 'постановл.*гибдд', 'отмен.*постановл']):
        return True
    
    return False


def is_moral_damage_allowed(template_data: dict) -> bool:
    """
    Определяет, разрешена ли компенсация морального вреда в данном иске.
    Запрещена в: наследстве, выселении, деликтах (затопление).
    """
    requirements = (template_data.get('requirements', '') or '').lower()
    situation = (template_data.get('situation', '') or '').lower()
    legal_basis = (template_data.get('legal_basis', '') or '').lower()
    
    combined = f"{requirements} {situation} {legal_basis}"
    
    # Признаки категорий, где моральный вред ЗАПРЕЩЁН
    forbidden_markers = [
        'выселить',
        'утративш.*прав',
        'восстановить срок.*наследств',
        'установить факт.*наследств',
        'затоп',
        'залит',
        '1064.*гк',  # деликт
        '31.*жк',  # выселение
        '35.*жк',  # выселение
    ]
    
    for marker in forbidden_markers:
        if re.search(marker, combined):
            return False
    
    return True


def split_into_items(text: str) -> list:
    """Разбивает текст на пункты с удалением дубликатов и мусора"""
    if not text:
        return []
    
    text = remove_duplicate_headers(text)
    
    lines = []
    lines_by_newline = [line.strip() for line in text.split('\n') if line.strip()]
    
    if len(lines_by_newline) > 1:
        lines = lines_by_newline
    else:
        lines = re.split(r'\s+(?=\d+[\.\)])\s*', text)
        lines = [line.strip() for line in lines if line.strip()]
    
    if len(lines) <= 1 and ';' in text:
        lines = [line.strip() for line in text.split(';') if line.strip()]
    
    if len(lines) == 0:
        lines = [text]
    
    merged_lines = []
    for line in lines:
        starts_with_number = re.match(r'^\d+[\.\)]\s*', line)
        
        if starts_with_number:
            merged_lines.append(line)
        elif len(line) < 15 and merged_lines:
            merged_lines[-1] = merged_lines[-1] + ' ' + line
        else:
            merged_lines.append(line)
    
    unique_lines = []
    seen = set()
    for line in merged_lines:
        line_without_number = re.sub(r'^\d+[\.\)]\s*', '', line).strip()
        line_cleaned = clean_template_hints(line_without_number).lower()
        
        if any(x in line_cleaned for x in ['прошу суд', 'требую', 'прошу', 'ходатайствую']):
            continue
        
        if any(x in line_cleaned for x in ['не заявлены', 'отсутствуют', 'не подлежат', 'ввиду отсутствия']):
            continue
        
        if line_cleaned not in seen and line_cleaned:
            seen.add(line_cleaned)
            cleaned_line = clean_template_hints(line)
            if cleaned_line:
                unique_lines.append(cleaned_line)
    
    return unique_lines


def fix_unclosed_quotes(text: str) -> str:
    """Исправляет незакрытые кавычки"""
    if not text:
        return text
    
    open_quotes = text.count('«')
    close_quotes = text.count('»')
    
    if open_quotes > close_quotes:
        text = text + '»' * (open_quotes - close_quotes)
    
    return text


def ensure_jurisdiction(recipient: str, template_data: dict = None) -> str:
    """
    Для исков в суд определяет правильную подсудность на основе типа спора.
    - ЗоЗПП → по месту жительства истца (ст. 17 ЗоЗПП)
    - Недвижимость → по месту нахождения имущества (ст. 30 ГПК РФ)
    - Деликты (затопление) → по месту жительства ответчика (ст. 28 ГПК РФ)
    - По умолчанию → по месту жительства ответчика
    """
    if not recipient:
        return "В районный суд по месту жительства ответчика"
    
    recipient_lower = recipient.lower()
    
    # Если подсудность уже указана — не трогаем
    if 'по месту жительства' in recipient_lower or \
       'по месту нахождения' in recipient_lower or \
       'мировой' in recipient_lower or \
       'арбитраж' in recipient_lower:
        return recipient
    
    # Если это суд, но нет указания подсудности
    if 'суд' in recipient_lower:
        situation = ''
        requirements = ''
        legal_basis = ''
        if template_data:
            situation = (template_data.get('situation', '') or '').lower()
            requirements = (template_data.get('requirements', '') or '').lower()
            legal_basis = (template_data.get('legal_basis', '') or '').lower()
        
        combined = f"{situation} {requirements} {legal_basis}"
        
        # Для ЗоЗПП — по месту жительства истца (ст. 17 ЗоЗПП)
        if any(x in combined for x in ['потребител', 'зозпп', 'магазин', 'товар', 'неустойк']):
            recipient = recipient + ' по месту жительства истца'
        # Для недвижимости — по месту нахождения (ст. 30 ГПК РФ)
        elif any(x in combined for x in ['квартир', 'дом', 'земельн', 'недвижим', 
                                          'высел', 'утративш.*прав', 'наследств',
                                          'признан.*прав.*собственн', 'раздел имуществ']):
            recipient = recipient + ' по месту нахождения недвижимого имущества'
        # Для деликтов (затопление) — по месту жительства ответчика (ст. 28 ГПК РФ)
        elif any(x in combined for x in ['затоп', 'залит', '1064', 'ущерб.*сосед']):
            recipient = recipient + ' по месту жительства ответчика'
        # По умолчанию — по месту жительства ответчика (ст. 28 ГПК РФ)
        else:
            recipient = recipient + ' по месту жительства ответчика'
    
    return recipient


def generate_legal_document(template_data: dict, output_dir: str, doc_type: str = "претензия") -> str:
    """Создаёт Word-документ на основе данных шаблона"""
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # === ШАПКА ===
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    recipient = template_data.get('recipient', '')
    recipient = clean_template_hints(recipient)
    
    # Для иска — обеспечиваем правильную подсудность
    if doc_type == "иск":
        recipient = ensure_jurisdiction(recipient, template_data)
    
    if not recipient:
        recipient = '[Наименование]'
    run = p.add_run(recipient)
    run.font.size = Pt(11)
    
    if doc_type == "иск":
        defendant = template_data.get('defendant', '').strip()
        defendant = clean_template_hints(defendant)
        if not defendant and recipient and recipient != '[Наименование]' and 'суд' not in recipient.lower():
            defendant = recipient
        if defendant:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(f"Ответчик: {defendant}")
            run.font.size = Pt(11)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sender = template_data.get('sender', '')
    sender = clean_template_hints(sender)
    
    if doc_type == "иск":
        if sender:
            run = p.add_run(f"Истец: {sender}")
        else:
            run = p.add_run("Истец: ___________________________________\n(ФИО, адрес, телефон, паспорт)")
    else:
        if sender:
            run = p.add_run(f"от {sender}")
        else:
            run = p.add_run("от ___________________________________\n(ФИО, адрес, телефон)")
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    # === ЗАГОЛОВОК ===
    doc_titles = {
        "претензия": "ПРЕТЕНЗИЯ",
        "жалоба": "ЖАЛОБА",
        "иск": "ИСКОВОЕ ЗАЯВЛЕНИЕ",
        "ходатайство": "ХОДАТАЙСТВО"
    }
    doc_title = doc_titles.get(doc_type, doc_type.upper())
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(doc_title)
    run.bold = True
    run.font.size = Pt(14)
    
    doc.add_paragraph()
    
    # === СИТУАЦИЯ ===
    situation = template_data.get('situation', '')
    situation = clean_template_hints(situation)
    situation = fix_unclosed_quotes(situation)
    if not situation:
        situation = '[Описание ситуации]'
    doc.add_paragraph(situation)
    doc.add_paragraph()
    
    # === ПРАВОВОЕ ОБОСНОВАНИЕ ===
    legal_basis = template_data.get('legal_basis', '').strip()
    legal_basis = clean_template_hints(legal_basis)
    legal_basis = fix_unclosed_quotes(legal_basis)
    if not legal_basis:
        legal_basis = ("В соответствии с действующим законодательством Российской Федерации, "
                      "а также на основании прав, предоставленных мне нормативными правовыми актами, "
                      "регулирующими данную ситуацию.")
    
    p = doc.add_paragraph()
    run = p.add_run("Правовое обоснование: ")
    run.bold = True
    p.add_run(legal_basis)
    doc.add_paragraph()
    
    # === ЦЕНА ИСКА ===
    if doc_type == "иск":
        claim_amount = template_data.get('claim_amount', '').strip()
        claim_amount = clean_template_hints(claim_amount)
        claim_amount = remove_duplicate_headers(claim_amount)
        claim_amount = format_claim_amount(claim_amount)
        if claim_amount:
            p = doc.add_paragraph()
            run = p.add_run("Цена иска: ")
            run.bold = True
            p.add_run(claim_amount)
            doc.add_paragraph()
    
    # === ТРЕБОВАНИЯ ===
    req_titles = {
        "претензия": "На основании вышеизложенного ТРЕБУЮ:",
        "жалоба": "На основании вышеизложенного ПРОШУ:",
        "иск": "ПРОШУ СУД:",
        "ходатайство": "На основании вышеизложенного ХОДАТАЙСТВУЮ:"
    }
    req_title = req_titles.get(doc_type, "ТРЕБУЮ:")
    
    p = doc.add_paragraph()
    run = p.add_run(req_title)
    run.bold = True
    
    requirements = template_data.get('requirements', '').strip()
    requirements = clean_template_hints(requirements)
    requirements = remove_duplicate_headers(requirements)
    requirements = fix_unclosed_quotes(requirements)
    
    # === УДАЛЯЕМ МОРЛЬНЫЙ ВРЕД, ЕСЛИ ОН НЕ РАЗРЕШЁН ===
    if doc_type == "иск" and not is_moral_damage_allowed(template_data):
        # Удаляем пункты о моральном вреде из требований
        req_lines_temp = split_into_items(requirements)
        filtered_lines = []
        for line in req_lines_temp:
            line_lower = line.lower()
            if 'моральн' in line_lower and ('вред' in line_lower or 'ущерб' in line_lower):
                continue
            filtered_lines.append(line)
        requirements = '\n'.join(filtered_lines)
    
    if not requirements:
        requirements = "Удовлетворить мои законные требования в соответствии с действующим законодательством РФ."
    
    req_lines = split_into_items(requirements)
    
    for i, line in enumerate(req_lines, 1):
        line = re.sub(r'^\d+[\.\)]\s*', '', line)
        line = clean_template_hints(line)
        line = fix_unclosed_quotes(line)
        if line:
            p = doc.add_paragraph(f"{i}. {line}")
            p.paragraph_format.left_indent = Pt(20)
    
    doc.add_paragraph()
    
    # === ПРИЛОЖЕНИЯ ===
    p = doc.add_paragraph()
    run = p.add_run("Приложения:")
    run.bold = True
    
    attachments = template_data.get('attachments', '').strip()
    attachments = clean_template_hints(attachments)
    if not attachments:
        attachments = "1. Копия данного документа\n2. Копии документов, подтверждающих требования"
    
    att_lines = split_into_items(attachments)
    
    if doc_type == "иск":
        att_text = ' '.join(att_lines).lower()
        
        # === ГОСПОШЛИНА ===
        fee_exempt = is_fee_exempt_category(template_data)
        
        if fee_exempt:
            att_lines = [line for line in att_lines 
                         if 'госпошлин' not in line.lower() and 'государственной пошлин' not in line.lower()]
        else:
            if 'госпошлин' not in att_text and 'государственной пошлин' not in att_text:
                att_lines.append("Квитанция об уплате государственной пошлины")
        
        # Удаляем мусорные приложения
        att_lines = [line for line in att_lines if not is_garbage_attachment(line)]
        
        # Копия иска для ответчика — всегда нужна
        if 'копия искового' not in att_text and 'для ответчика' not in att_text:
            att_lines.append("Копия искового заявления для ответчика")
        
        # === РАСЧЁТ ЦЕНЫ ИСКА — только для имущественных исков ===
        is_non_property = is_non_property_claim(template_data)
        
        if not is_non_property:
            if 'расчёт' not in att_text and 'расчет' not in att_text:
                att_lines.append("Расчёт цены иска")
        else:
            att_lines = [line for line in att_lines 
                         if 'расчёт цены' not in line.lower() and 'расчет цены' not in line.lower()]
    
    for i, line in enumerate(att_lines, 1):
        line = re.sub(r'^\d+[\.\)]\s*', '', line)
        line = clean_template_hints(line)
        if line:
            p = doc.add_paragraph(f"{i}. {line}")
            p.paragraph_format.left_indent = Pt(20)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # === ДАТА И ПОДПИСЬ ===
    p = doc.add_paragraph()
    p.add_run("Дата: ___________                    Подпись: ___________")
    
    # === СОХРАНЕНИЕ ===
    safe_name = re.sub(r'[^\w\-.а-яА-Я]', '_', doc_title)
    output_path = Path(output_dir) / f"{safe_name}.docx"
    doc.save(str(output_path))
    
    return str(output_path)
